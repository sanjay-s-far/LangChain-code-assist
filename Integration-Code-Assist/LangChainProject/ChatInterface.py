import asyncio
from pathlib import Path
import gradio as gr
from ContextTemplateBuilder import builder
from docx import Document
import PyPDF2
import os
import re
import zipfile
import time
# --- File paths ---
preview_doc = Path(r"C:\Users\sd116131\LangChain-code-assist\Integration-Code-Assist\FinishedFiles")


# --- Extract key-value placeholders from BRD ---
def extract_placeholders(brd_text):
    pattern = re.compile(r'^(.*?):\s+(.*)$', re.MULTILINE)
    mapping = {}
    for key, value in pattern.findall(brd_text):
        norm_key = normalize_key(key)
        mapping[norm_key] = value.strip()
    return mapping


def normalize_key(key):
    key = key.lower()
    key = key.replace(" ", "_").replace("-", "_")
    key = re.sub(r"[^\w_]", "", key)
    return key


def fill_placeholders(text, mapping):
    for key, value in mapping.items():
        # Replace various placeholder styles
        text = re.sub(rf"\b{key}\b", value, text, flags=re.IGNORECASE)
        text = re.sub(rf"\{{\{{{key}\}}\}}", value, text, flags=re.IGNORECASE)
        text = re.sub(rf"\${{\s*{key}\s*}}", value, text, flags=re.IGNORECASE)
        text = re.sub(rf"<<{key}>>", value, text, flags=re.IGNORECASE)
        text = re.sub(rf"\byour[-_]{key}\b", value, text, flags=re.IGNORECASE)
    return text


# --- Extract text from file ---
def extract_text_from_file(file):
    if file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        return "".join([page.extract_text() or "" for page in reader.pages]).strip()
    elif file.name.endswith(".docx"):
        doc = Document(file.name)
        return "\n".join(p.text for p in doc.paragraphs).strip()
    else:
        return "Unsupported file format."


# --- Read preview content ---
def retrieve_latest_preview_file(file_name):
    if not file_name.endswith(".txt"):
        file_name += ".txt"
    matching_files = [f for f in preview_doc.iterdir() if f.name == file_name]
    if not matching_files:
        return "No preview file found."
    with open(matching_files[0], "r", encoding="utf-8") as f:
        return f.read()


# --- Async builder ---
def generate_preview_from_requirement(user_msg, project_name):
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        return loop.run_until_complete(builder(user_msg, project_name))
    except Exception as e:
        return f"Error occurred: {str(e)}"


# --- Helpers ---
def clean_path(path: str) -> str:
    return re.sub(r"[\"']", "", path.strip())


def generate_files_and_zip(project_name: str, project_base_dir: str, txt_file_path: str, brd_mapping: dict) -> str:
    with open(txt_file_path, "r", encoding="utf-8") as f:
        text = f.read()

    # Replace placeholders in the file content
    filled_text = fill_placeholders(text, brd_mapping)

    # Extract file/code blocks
    pattern = re.compile(r"\*\*\d+\.\s+([^*]+)\*\*.*?\n(?:\w+)?\n(.*?)\n", re.DOTALL)
    matches = pattern.findall(filled_text)
    if not matches:
        raise ValueError("No file/code blocks matched the expected pattern.")

    for file_path_raw, code in matches:
        file_path = clean_path(file_path_raw)
        full_path = os.path.join(project_base_dir, file_path)
        os.makedirs(os.path.dirname(full_path), exist_ok=True)
        with open(full_path, "w", encoding="utf-8") as f:
            f.write(code.strip())

    zip_file_path = os.path.join(os.path.dirname(project_base_dir), f"{project_name}.zip")
    with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for foldername, _, filenames in os.walk(project_base_dir):
            for filename in filenames:
                filepath = os.path.join(foldername, filename)
                relpath = os.path.relpath(filepath, os.path.dirname(project_base_dir))
                zipf.write(filepath, relpath)

    return zip_file_path


# --- Save & regenerate after edit ---
def handle_brd_edit_save(edited_text, project_name):
    file_path = os.path.join(preview_doc, f"{project_name}.docx")
    doc = Document()
    for line in edited_text.splitlines():
        doc.add_paragraph(line)
    doc.save(file_path)

    brd_mapping = extract_placeholders(edited_text)
    generate_preview_from_requirement(edited_text, project_name)
    preview = retrieve_latest_preview_file(f"{project_name}.txt")

    project_txt_file = os.path.join(preview_doc, f"{project_name}.txt")
    generated_project_folder = os.path.join("generated_projects", project_name)
    os.makedirs(generated_project_folder, exist_ok=True)
    zip_file_path = generate_files_and_zip(project_name, generated_project_folder, project_txt_file, brd_mapping)

    return [[project_name, preview]], gr.update(value=zip_file_path)


# --- Gradio UI ---
with gr.Blocks() as demo:
    gr.Markdown("## 🧰 Spring Boot Project Generator ")

    with gr.Row():
        with gr.Column(scale=1):
            project_name_box = gr.Textbox(label="Project Name")
            file_upload = gr.File(label="Upload BRD (PDF/DOCX)", file_types=[".pdf", ".docx"])
            submit_btn = gr.Button("Generate Project")
            save_brd_btn = gr.Button("Save Edited BRD & Regenerate", visible=False)

        with gr.Column(scale=2):
            chatbot = gr.Chatbot(label="Generated Preview")
            download_btn = gr.File(label="Download ZIP", interactive=True)
            brd_content_box = gr.Textbox(label="Edit BRD Content", lines=20, interactive=True, visible=False)
            clear = gr.Button("Clear")


    # --- Handle form submission ---
    def handle_form_submission(project_name, file, history):
        if not project_name or not file:
            return history + [["Error", "Please provide both Project Name and a BRD file."]], None, "", gr.update(
                visible=False), gr.update(visible=False)

        for _ in range(5):
            if os.path.exists(file.name) and os.path.getsize(file.name) > 0:
                break
            time.sleep(0.5)

        text = extract_text_from_file(file)
        if not text or text.startswith("Unsupported"):
            return history + [["Error", text or "Could not extract text from file."]], None, "", gr.update(
                visible=False), gr.update(visible=False)

        brd_mapping = extract_placeholders(text)
        generate_preview_from_requirement(text, project_name)
        preview = retrieve_latest_preview_file(f"{project_name}.txt")
        filled_preview = fill_placeholders(preview, brd_mapping)

        # Save BRD
        file_path = os.path.join(preview_doc, f"{project_name}.docx")
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.save(file_path)

        # Generate ZIP
        project_txt_file = os.path.join(preview_doc, f"{project_name}.txt")
        generated_project_folder = os.path.join("generated_projects", project_name)
        os.makedirs(generated_project_folder, exist_ok=True)
        zip_file_path = generate_files_and_zip(project_name, generated_project_folder, project_txt_file, brd_mapping)

        return (
            history + [[project_name, filled_preview]],
            gr.update(value=zip_file_path),
            text,
            gr.update(visible=True),
            gr.update(visible=True)
        )


    # --- Event bindings ---
    submit_btn.click(
        handle_form_submission,
        inputs=[project_name_box, file_upload, chatbot],
        outputs=[chatbot, download_btn, brd_content_box, brd_content_box, save_brd_btn],
    )

    save_brd_btn.click(
        handle_brd_edit_save,
        inputs=[brd_content_box, project_name_box],
        outputs=[chatbot, download_btn],
    )

    clear.click(lambda: [], None, chatbot, queue=False)

demo.launch()
