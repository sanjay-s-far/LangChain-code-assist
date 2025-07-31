import os
import re
import zipfile
import asyncio
from pathlib import Path

import uvicorn
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, UploadFile, Form, File

from docx import Document
import PyPDF2
from ContextTemplateBuilder import builder


# === Config ===
project_dir = Path("FinishedFiles")
generated_dir = Path("LangChainProject/generated_projects")
project_dir.mkdir(exist_ok=True)
generated_dir.mkdir(parents=True, exist_ok=True)

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Restrict in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def normalize_key(key: str) -> str:
    key = key.lower().replace(" ", "_").replace("-", "_")
    return re.sub(r"[^\w_]", "", key)

def extract_placeholders(brd_text: str) -> dict:
    pattern = re.compile(r'^(.*?):\s+(.*)$', re.MULTILINE)
    return {normalize_key(k): v.strip() for k, v in pattern.findall(brd_text)}

def fill_placeholders(text: str, mapping: dict) -> str:
    for key, value in mapping.items():
        text = re.sub(rf"\b{key}\b", value, text, flags=re.IGNORECASE)
        text = re.sub(rf"\{{\{{{key}\}}\}}", value, text, flags=re.IGNORECASE)
        text = re.sub(rf"\${{\s*{key}\s*}}", value, text, flags=re.IGNORECASE)
        text = re.sub(rf"<<{key}>>", value, text, flags=re.IGNORECASE)
        text = re.sub(rf"\byour[-_]{key}\b", value, text, flags=re.IGNORECASE)
    return text

def extract_text_from_file(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "".join([page.extract_text() or "" for page in reader.pages])
    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    return "Unsupported file format."

def clean_path(path: str) -> str:
    return re.sub(r"[\"']", "", path.strip())

def generate_files_and_zip(project_name: str, base_dir: str, txt_file_path: str, brd_mapping: dict) -> str:
    with open(txt_file_path, "r", encoding="utf-8") as f:
        text = f.read()

    filled_text = fill_placeholders(text, brd_mapping)
    pattern = re.compile(r"\*\*\d+\.\s+([^*]+)\*\*.*?\n(?:\w+)?\n(.*?)\n", re.DOTALL)
    matches = pattern.findall(filled_text)

    if not matches:
        raise ValueError("No file/code blocks matched.")

    for file_path_raw, code in matches:
        file_path = clean_path(file_path_raw)
        full_path = os.path.join(base_dir, file_path)
        os.makedirs(os.path.dirname(full_path), exist_ok=True)
        with open(full_path, "w", encoding="utf-8") as f:
            f.write(code.strip())

    zip_file_path = os.path.join(generated_dir, f"{project_name}.zip")
    with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for foldername, _, filenames in os.walk(base_dir):
            for filename in filenames:
                filepath = os.path.join(foldername, filename)
                relpath = os.path.relpath(filepath, generated_dir)
                zipf.write(filepath, relpath)

    return zip_file_path

@app.post("/generate")
async def generate_project(project_name: str = Form(...), file: UploadFile = File(...)):
    print(f"[INFO] Received project_name: {project_name}, file: {file.filename}")

    file_path = project_dir / file.filename
    with open(file_path, "wb") as f:
        f.write(await file.read())

    text = extract_text_from_file(str(file_path))
    if not text or text.startswith("Unsupported"):
        return JSONResponse(status_code=400, content={"error": "Unsupported or empty file"})

    brd_mapping = extract_placeholders(text)

    # Run builder
    try:
        await asyncio.to_thread(builder, text, project_name)
    except Exception as e:
        print(f"[ERROR] Builder failed: {e}")
        return JSONResponse(status_code=500, content={"error": "Builder failed."})

    preview_file = project_dir / f"{project_name}.txt"
    if not preview_file.exists():
        return JSONResponse(status_code=500, content={"error": "Preview file not found."})

    filled_text = fill_placeholders(preview_file.read_text(), brd_mapping)

    # Save BRD as docx
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(project_dir / f"{project_name}.docx")

    # Generate final project
    output_folder = generated_dir / project_name
    output_folder.mkdir(exist_ok=True)
    zip_path = generate_files_and_zip(project_name, str(output_folder), str(preview_file), brd_mapping)

    return JSONResponse(status_code=200, content={"project_name": project_name, "zip_path": zip_path})

@app.get("/")
def root():
    return {"message": "Spring Boot Project Generator is running!"}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
